# процедура для копирования поля из одной модели в другую
# синхронизирует поле Описание в двух моделях


import SIMMA_Class as All_Class
from docx import Document
from docx.shared import Pt, RGBColor
import html
import json
import  requests
import docraptor
import sys
from PyQt5 import QtWidgets, uic
import pyqtgraph as pg

sid="6a86cba0afd546639283fa16aebd9171"
mid1="m1f7"
mid2="m1d3"
class_1="oe5"
class_2="o16b1"
list_atribute_1=[]
list_atribute__2=[]
list_element_1=[]
list_element_2=[]
url="https://metalogica.ru"

def edit_element():
    sc = All_Class.Simma_class
    sc.o_id = class_1
    list_atribute_1 = sc.gca(sc, url, mid1, sid)
    list_for_sevt = ["descr"]
    #for tmp in list_atribute_1:
    #   list_for_sevt.append(tmp[0])
    sc.sevt(sc, url, mid1, sid, list_for_sevt, [])
    list_element_1 = sc.gtvt(sc, url, mid1, sid)
    sc.sevt(sc, url, mid1, sid, ['name'], [])

    sc2 = All_Class.Simma_class
    sc2.o_id = class_2
    list_atribute_2 = sc2.gca(sc2, url, mid2, sid)
    list_for_sevt = ["descr","a4dad"]
    #for tmp in list_atribute_2:
    #    list_for_sevt.append(tmp[0])
    sc2.sevt(sc2, url, mid2, sid, list_for_sevt, [])
    list_element_2 = sc2.gtvt(sc2, url, mid2, sid)
    sc2.sevt(sc2, url, mid2, sid, ['name'], [])

    for tmp_element_1 in list_element_1:
        for tmp_element_2 in list_element_2:
            if tmp_element_1[1]==tmp_element_2[3][1]["f2"]:
               s_el1 = All_Class.Simma_element
               s_el2=All_Class.Simma_element
               s_el1.e_id=tmp_element_1[0]
               about_element=s_el1.gea(s_el1,url,mid1,sid)
               s_el2.e_id=tmp_element_2[0]
               s_el2.ee(s_el2,url,mid2,sid,class_2,["descr"],[about_element[3][1]])
    pass

def test_word():

    # создание документа
    doc = Document()
    # задаем стиль текста по умолчанию
    style = doc.styles['Normal']
    # название шрифта
    style.font.name = 'Calibri'
    # размер шрифта
    style.font.size = Pt(14)
    p = doc.add_paragraph('Пользовательское ')
    # добавляем текст прогоном
    run = p.add_run('форматирование ')
    # размер шрифта
    run.font.size = Pt(16)
    # курсив
    run.font.italic = True
    # добавляем еще текст прогоном
    run = p.add_run('символов текста.')
    # Форматируем:
    # название шрифта
    run.font.name = 'Arial'
    # размер шрифта
    run.font.size = Pt(18)
    # цвет текста
    run.font.color.rgb = RGBColor(255, 0, 0)
    # + жирный и подчеркнутый
    run.font.bold = True
    run.font.underline = True
    doc.save('test.docx')


def test_html():

    doc, tag, text = Doc().tagtext()
    element=All_Class.Simma_element
    element.e_id="e83f97"
    about_element=element.gea(element,url,mid1,sid)
    str_line=""
    rezult=about_element[10][1]
    soup = html.unescape(rezult)
    file = open("Test.html", "w")
    file.write("<html>")
    file.write("<head>")
    file.write("<title>My Webpage</title>")
    file.write("</head>")
    file.write("<body>")
    file.write("<h1>Welcome to my webpage!</h1>")
    file.write("</body>")
    file.write(soup)
    file.close()

class MainWindow(QtWidgets.QMainWindow):
    def __init__(self, *args, **kwargs):
        super(MainWindow, self).__init__(*args, **kwargs)

        # Загрузите страницу интерфейса
        uic.loadUi('Grafik.ui', self)
        grid = QtWidgets.QGridLayout(self.Graphwidget)
        grid.addWidget(self.Graphwidget, 0, 0)

        self.plot([1,2,3,4,5,6,7,8,9,10], [30,32,34,32,33,31,29,32,35,10])

    # мы добавили метод plot(), который принимает два массива:
    # temperature и hour, затем строит данные с помощью метода graphWidget.plot().

    def plot(self, hour, temperature):
        self.Graphwidget.plot(hour, temperature)

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    main = MainWindow()
    main.show()
    sys.exit(app.exec_())

