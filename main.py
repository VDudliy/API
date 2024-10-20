import sys
from PyQt5 import QtWidgets, QtCore
from lxml.doctestcompare import strip

import SIMMA_Class as All_Class
from PyQt5.uic import loadUi
from PyQt5.QtCore import Qt, QEvent
from PyQt5.QtWidgets import QTreeWidgetItem, QMessageBox, QApplication
from PyQt5.QtChart import QChart, QChartView,QLineSeries,QPieSeries,QPieSlice,QBarSet,QStackedBarSeries,QBarCategoryAxis, QValueAxis

from PyQt5.QtGui import QPen
import json
import copy
import docx
from bs4 import BeautifulSoup
from docx.shared import Pt, RGBColor,Mm
from datetime import date
import re
from PyQt5.QtCore import QPoint

class ui_Simma():
    url = ""
    mid = "m1f7"
    login = ""
    password = ""
    sid = ""
    dict_class={}
    dict_element_level_1= {}  # элементы класса "o1562":"Отчет"
    dict_element_level_2 = {} # элементы класса "o1563":"Отчет раздела"
    dict_element_level_3 = {} # элементы класса "o149b,":"API-метод"
    dict_atribut_level_1= {} # атрибуты класса "o1562":"Отчет"
    dict_atribut_level_2= {} # атрибуты класса "o1563":"Отчет раздела"
    dict_atribut_level_3= {} # атрибуты класса "o149b,":"API-метод"
    id_attr_index="a53f4"  # атрибут содержащий поле с индексом для отчета
    id_attr_class_for_content="a53f3"  # атрибут с именем класса для которого предназначен отчет
    id_attr_colontitul_list_1_up="a53fb" # верхний колонтитул первого листа
    id_attr_colontitul_list_1_down = "a53fc" # нижний колонтитул первого листа
    id_attr_colontitul_list_all_up = "a53fd" # верхний колонтитул всех листов
    id_attr_colontitul_list_all_down = "a53fe" # нижний колонтитул всех листов
    id_attr_text_list_1 = "a53fa" # атрибут с текстом первого листа
    id_attr_level2_preambula="a5401" # атрибут с преамбулой у разделов
    id_attr_section="a53ff" # атрибут в котором описано -какие разделы содержит отчет. Это связь.
    id_attr_status="a53b7" # атрибут статуса в каталлоге контента
    dict_element_index={}
    dict_element_index3 = {}
    dict_class_content={}
    dict_item_content={}
    dict_list_level1_2={}
    dict_attribute_config={}# сюда считаем из конфигурацйии атрибуты
    list_level_element=[]
    parametr=0

    def __init__(self):
        self.ui_Simma = loadUi("Simma_API.ui")
        self.set_log_data()
        self.ui_Simma.tabWidget_Auth.currentChanged.connect(self.set_type_report)
        self.ui_Simma.pushButton_Auth.clicked.connect(self.auth)
        self.ui_Simma.treeWidget_SIMMA.setContextMenuPolicy(Qt.CustomContextMenu)
        self.ui_Simma.treeWidget_SIMMA.customContextMenuRequested.connect(self.contextMenuEvent)
        self.ui_Simma.comboBox_report.currentTextChanged.connect(self.set_type_report)
        self.ui_Simma.pushButton_all.clicked.connect(self.all_doc_html)
        self.ui_Simma.pushButton_exit.clicked.connect(self.exit_sys)
        self.ui_Simma.pushButton_slice.clicked.connect(self.slice_doc_html)
        self.ui_Simma.treeWidget_SIMMA.itemClicked.connect(self.item_view)
        self.ui_Simma.pushButton_create_diagramm.clicked.connect(self.def_diagramm)
        self.ui_Simma.show()
    def set_log_data(self):
        try:
            self.show_logo()
            with open("config_loader_json.txt", 'r') as json_file:
                data = json.load(json_file)
                for txt in data:
                    if txt == 'url':
                        self.url = data[txt]
                    elif txt == 'login':
                        self.login = data[txt]
                    elif txt =="dict_class":
                        self.dict_class = data[txt]
                    elif txt=="dict_attribute":
                        self.dict_attribute_config=data[txt]
                for k,v in  self.dict_attribute_config.items():
                    match k:
                         case "id_attr_index":
                               self.id_attr_index=v
                         case "id_attr_class_for_content":
                               self.id_attr_class_for_content=v
                         case "id_attr_colontitul_list_1_up":
                               self.id_attr_colontitul_list_1_up=v
                         case "id_attr_colontitul_list_1_down":
                               self.id_attr_colontitul_list_1_down=v
                         case "id_attr_colontitul_list_all_up":
                               self.id_attr_colontitul_list_all_up=v
                         case "id_attr_colontitul_list_all_down":
                               self.id_attr_colontitul_list_all_up=v
                         case "id_attr_text_list_1":
                               self.id_attr_text_list_1=v
                         case "id_attr_level2_preambula":
                               self.id_attr_level2_preambula=v
                         case "id_attr_section":
                               self.id_attr_section=v
        except:
            self.url = 'https://www.metalogica.ru/'
            self.login = 'vadim_d'
        self.ui_Simma.lineEdit_URL.setText(self.url)
        self.ui_Simma.lineEdit_login.setText(self.login)
    def auth(self):
        if self.sid == "":
            lg = All_Class.Simma_application
            self.sid = lg.logining(lg, self.url, self.ui_Simma.lineEdit_login.text(),
                                   self.ui_Simma.lineEdit_password.text())
            if self.sid['state'] == 'ok':
                self.sid = self.sid['sid']
                self.ui_Simma.label_4.setText("Авторизация прошла успешно")

                self.upload_data_from_simma()
                self.set_box_report()

            else:
                self.ui_Simma.label_4.setText("Ошибка авторизации. Повторите попытку")
            with open('config_loader_json.txt', 'w') as json_file:
                data = {'url': self.ui_Simma.lineEdit_URL.text(),
                        "login": self.ui_Simma.lineEdit_login.text(),
                        "dict_class": self.dict_class,
                        "dict_attribute":self.dict_attribute_config
                        }
                json.dump(data, json_file)
                json_file.close()
        elif len(self.sid) == 32:
            self.show_Ok_messagebox(2, "")
        else:
            self.sid = ""
    def exit_sys(self):
        sys.exit()
    def show_logo(self):
        file2 = open('logo.html', "r", encoding='utf-16')
        data=file2.read()
        self.ui_Simma.label_logo.setText(data)
        file2.close
    def show_critical_messagebox(self, param):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        if param == 1:
            # setting message for Message Box
            msg.setText("Не пройдена авторизация")
        elif param == 2:
            msg.setText("Не правильно выбрана связь. Связи элементов не созданы!")

        # setting Message box window title
        msg.setWindowTitle("Critical MessageBox")

        # declaring buttons on Message Box
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)

        # start the app
        retval = msg.exec_()
    def show_Ok_messagebox(self, number, prefix):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Information)
         # setting message for Message Box
        if number == 1:
            msg.setText("Документ сформирован! Размещен в файле: \"" + self.ui_Simma.comboBox_report.currentText() + prefix+"\"")
        elif number == 2:
            msg.setText("Вы уже авторизованы!")
        elif number==3:
            msg.setText("Документ сформирован! Размещен в файле: \"" + self.ui_Simma.comboBox_report.currentText() + "_"+ self.ui_Simma.treeWidget_SIMMA.currentItem().text(0)+prefix + "\"")
        # setting Message box window title
        msg.setWindowTitle("Information")

        # declaring buttons on Message Box
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)

        # start the app
        retval = msg.exec_()
    def set_box_report(self):
        if len(self.sid)==32:
           sc = All_Class.Simma_class
           sc.o_id=list(self.dict_class.keys())[0]
           about_class = sc.gtvt(sc, self.url, self.mid, self.sid)
           for tmp in about_class:
               self.ui_Simma.comboBox_report.addItem(tmp[1])
    def set_type_report(self):
        if len(self.sid) == 32:
           self.make_tree()
           self.init_tree()
           self.init_graph()
        elif self.ui_Simma.tabWidget_Auth.currentIndex()==1 or self.ui_Simma.tabWidget_Auth.currentIndex()==2:
           self.show_critical_messagebox(1)
           self.ui_Simma.tabWidget_Auth.setCurrentIndex(0)
    def upload_data_from_simma(self):
        # нужно определить какие атрибуты в каждом классе, установить возврат всех атрибутов и считать все элемекнты в каждом классе
        # последовательность получается такая -  gca -> sevt -> gtvt
        # после окончания считывыания - вернуть настройки обратно
        # данные получае только по 2-м классам из конфигурациооного файла, классы с контентом необходимо найти по полю "для каких классов подходит отчет"
        countBar = 0
        self.ui_Simma.progressBar.maximum = 100
        level=1
        for k,v in self.dict_class.items():
            if level==1:
                sc = All_Class.Simma_class
                sc.o_id = k
                about_class = sc.gca(sc, self.url, self.mid, self.sid)
                self.dict_atribut_level_1[k]=about_class
                list_for_sevt=[]
                for tmp in about_class:
                    list_for_sevt.append(tmp[0])
                list_for_sevt.append("descr")
                sc.sevt(sc,self.url,self.mid,self.sid,list_for_sevt,[])
                about_class=sc.gtvt(sc, self.url, self.mid, self.sid)
                list_report=[]
                for tmp in about_class:
                    self.dict_element_level_1[tmp[0]]=tmp
                    list_report.append(tmp[1])
                sc.sevt(sc, self.url, self.mid, self.sid, ['name'], [])
                level=2
                valBar = 20
                self.ui_Simma.progressBar.setValue(valBar)
            elif level == 2:
                sc = All_Class.Simma_class
                sc.o_id = k
                about_class = sc.gca(sc, self.url, self.mid, self.sid)
                self.dict_atribut_level_2[k] = about_class
                list_for_sevt = []
                for tmp in about_class:
                    list_for_sevt.append(tmp[0])
                list_for_sevt.append("descr")
                sc.sevt(sc, self.url, self.mid, self.sid, list_for_sevt, [])
                about_class = sc.gtvt(sc, self.url, self.mid, self.sid)
                dict_element_level_2={} #это внутренний промежуточный
                for tmp in about_class:
                    dict_element_level_2[tmp[0]] = tmp
                sc.sevt(sc, self.url, self.mid, self.sid, ['name'], [])
                for tmp in list_report:
                    tmp_dict = {}
                    for k,v in dict_element_level_2.items():
                        if v[3][0]["f3"][0][1] == tmp:
                           tmp_dict[k]=v
                    self.dict_element_level_2[tmp]=copy.deepcopy(tmp_dict)
                level = 3
                valBar = 40
                self.ui_Simma.progressBar.setValue(valBar)
        # здесь вставили процедуру определения классов и уже по нима с ключем наименования класса сформировать словарь третьега уровня
        list_class_content=[]
        for k, v in self.dict_element_level_1.items():
            for tmp in v[3]:
                    if tmp["f1"] == self.id_attr_class_for_content:
                        list_class_content.append(tmp["f2"])
                        self.dict_item_content[v[1]]=tmp["f2"]
        sc=All_Class.Simma_class
        about_model=sc.gc(sc,self.url,self.mid,self.sid)
        for tmp in about_model:
            for  tmp_class in list_class_content:
                 if tmp[1] == tmp_class:
                     self.dict_class_content[tmp_class]=tmp[0]
        valBar = 60
        self.ui_Simma.progressBar.setValue(valBar)
        for k,v in self.dict_class_content.items():
            sc = All_Class.Simma_class
            sc.o_id = v
            about_class = sc.gca(sc, self.url, self.mid, self.sid)
            self.dict_atribut_level_3[v] = about_class
            list_for_sevt = []
            for tmp in about_class:
                list_for_sevt.append(tmp[0])
            sc.sevt(sc, self.url, self.mid, self.sid, list_for_sevt, [])
            about_class = sc.gtvt(sc, self.url, self.mid, self.sid)
            tmp_dict={}
            for tmp in about_class:
                tmp_dict[tmp[0]] = tmp
            self.dict_element_level_3[k] = copy.deepcopy(tmp_dict)
            sc.sevt(sc, self.url, self.mid, self.sid, ['name'], [])
            # получили словарь элементов третьего уровня класс:список элементов (имя класса прописано у элемента первого уровня)
            valBar = 100
            self.ui_Simma.progressBar.setValue(valBar)
    def make_tree(self):
       # из считанных данных строим список первого уровня

       for key in self.dict_item_content.keys():
           dict_element_index = {}
           for k, v in self.dict_element_level_1.items():
                if v[1]==key:
                    list_level1=[]
                    for tmp in v[3]:
                        if tmp["f3"] is not None:
                            for tmp1 in tmp["f3"]:
                                tmp_element = self.dict_element_level_2[key][tmp1[2]]
                                for tmp_element1 in tmp_element[3]:
                                    if tmp_element1["f1"] == self.id_attr_index:
                                        len_index = len(tmp_element1["f2"])
                                        if len_index == 1:
                                            list_level1.append(tmp1[2])
                                            dict_element_index[tmp_element[0]]=tmp_element1["f2"]
                                        else:
                                            dict_element_index[tmp_element[0]] = tmp_element1["f2"]
           self.dict_element_index[key]=copy.deepcopy(dict_element_index)
           self.dict_list_level1_2[key] = list_level1
        # построим словарь элементов в индексами третьего уровня
       for key in list(self.dict_item_content.keys()):
           dict_element_index3 = {}
           for k, v in self.dict_element_level_3[self.dict_item_content[key]].items():
               for tmp in v[3]:
                   if tmp["f1"]==self.id_attr_index:
                      dict_element_index3[k]=tmp["f2"]
           self.dict_element_index3[key]=copy.deepcopy(dict_element_index3)
        # построили список с элементами верхнего уровня, теперь его нужно наполнить вторым уровнем
       for key in self.dict_item_content.keys():
           for  tmp_element in self.dict_list_level1_2[key]:
                 tmp_list = []
                 for k,v in self.dict_element_index[key].items():
                     if self.dict_element_index[key][tmp_element]==v[0] and len(v)>1:
                         tmp_list.append(k)
                 if  len(tmp_list)>0:
                     tmp_dict={}
                     tmp_dict[tmp_element]=tmp_list
                     self.dict_list_level1_2[key][self.dict_list_level1_2[key].index(tmp_element)]=copy.deepcopy(tmp_dict)
    def item_view(self):
        name_element=self.ui_Simma.treeWidget_SIMMA.currentItem().text(0)
        list_for_label=""
        list_attribute_name = self.dict_atribut_level_3[self.dict_class_content[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]]]
        for k,v in self.dict_element_level_3[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]].items():
            if v[1]==name_element:
                for tmp_attr in v[3]:
                    if tmp_attr["f4"] != "r" and tmp_attr["f2"] is not None and tmp_attr["f2"] != "":
                        for attr in list_attribute_name:
                            if tmp_attr["f1"] in attr:
                                if attr[4] == "s":
                                    list_for_label=list_for_label+("<p align=left><b>" + attr[1] + "</b> : " + tmp_attr[
                                            "f2"] + "</p>")
                                elif tmp_attr["f1"] != self.id_attr_index:
                                    list_for_label=list_for_label+("<h4 align=left>" + attr[1] + "</h4>")
                                    list_for_label=list_for_label+("<div align=left>" + tmp_attr["f2"] + "</div>")
        for key,val in self.dict_class.items():
            if val=="Раздел отчета":
               list_attribute_name = self.dict_atribut_level_2[key]
        for k, v in self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()].items():
            if v[1] == name_element:
                for tmp_attr in v[3]:
                    if tmp_attr["f1"] == self.id_attr_level2_preambula:
                       if tmp_attr["f2"] is not None and tmp_attr["f2"] !="":
                           list_for_label=list_for_label+("<h4 align=left> Преамбула раздела </h4>")
                           list_for_label=list_for_label+("<div align=left>" + tmp_attr["f2"] + "</div>")
        list_for_label="<H2>"+name_element+"</H2>" +list_for_label
        self.ui_Simma.textEdit_item.setText(list_for_label)
    def recursive(self,dict_level3,index_parent,level,sg):  # cловарь с индексами элементов 3-го уровня, индекс родителя и уровень вложенности
        element_chaild=""
        for k,v in dict_level3.items():
            if len(index_parent)==1 and v is not None:
                if v[:level]==index_parent and len(v)==level+2:
                   element_chaild=k
                   ss = QtWidgets.QTreeWidgetItem(sg)
                   ss.setText(0, self.dict_element_level_3[
                       self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]][k][1])
                   ss.addChild(sg)
                   if element_chaild != "":
                       self.recursive(dict_level3, v, level + 2, ss)
            elif len(index_parent)>1 and v is not None:
                if v[:level] == index_parent and len(v)==level+2:
                    element_chaild = k
                    ss = QtWidgets.QTreeWidgetItem(sg)
                    ss.setText(0, self.dict_element_level_3[
                        self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]][k][1])
                    ss.addChild(sg)
                    if element_chaild!="":
                       self.recursive(dict_level3,v,level+2,ss)
        return(element_chaild) # вернули id элемента дочернего
    def recursive_for_report_docx(self,dict_level3,index_parent,level,doc):
        element_chaild = ""
        for k, v in dict_level3.items():
            if len(index_parent) == 1 and v is not None:
                if v[:level] == index_parent and len(v) == level + 2:
                    element_chaild = k
                    #выводим элемент к
                    p1 = doc.add_paragraph()
                    #непосредственно вывод
                    run = p1.add_run(k)
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    fmt = p1.paragraph_format
                    fmt.first_line_indent = Mm(15)
                    if element_chaild != "":
                        self.recursive_for_report(dict_level3, v, level + 2, doc)
            elif len(index_parent) > 1 and v is not None:
                if v[:level] == index_parent and len(v) == level + 2:
                    element_chaild = k
                    # выводим элемент к
                    p1 = doc.add_paragraph()
                    run = p1.add_run(k)
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    fmt = p1.paragraph_format
                    fmt.first_line_indent = Mm(15)
                    if element_chaild != "":
                        self.recursive_for_report(dict_level3, v, level + 2, doc)
        return (element_chaild)  # вернули id элемента дочернего
    def recursive_for_report_html(self,dict_level3,index_parent,level,file):
        element_chaild = ""
        for k, v in dict_level3.items():
            if len(index_parent) == 1 and v is not None:
                if v[:level] == index_parent and len(v) == level + 2:
                    element_chaild = k
                    #выводим элемент к
                    report=self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]
                    element=self.dict_element_level_3[report][element_chaild]
                    element_index=self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()][element_chaild]
                    list_attribute_name=self.dict_atribut_level_3[self.dict_class_content[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]]]
                    file.write("<h3 align=left style=margin-left:70 id="+element_chaild+">" + element_index+" "+element[1]+ "</h3>")
                    for tmp_attr in element[3]:
                        if tmp_attr["f4"]!="r" and tmp_attr["f2"] is not None and tmp_attr["f2"]!="":
                           for attr in list_attribute_name:
                               if tmp_attr["f1"] in attr:
                                  if attr[4]=="s":
                                       file.write("<p align=left style=margin-left:90><b>" + attr[1] + "</b> : " + tmp_attr["f2"] + "</p>")
                                  elif tmp_attr["f1"]!=self.id_attr_index:
                                       file.write("<h4 align=left style=margin-left:90>" + attr[1] + ": </h4>")
                                       file.write("<div align=left style=margin-left:90>" + tmp_attr["f2"] + "</div>")
                    #непосредственно вывод
                    if element_chaild != "":
                        self.recursive_for_report_html(dict_level3, v, level + 2, file)
            elif len(index_parent) > 1 and v is not None:
                if v[:level] == index_parent and len(v) == level + 2:
                    element_chaild = k
                    # выводим элемент к
                    report = self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]
                    element = self.dict_element_level_3[report][element_chaild]
                    element_index = self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()][
                        element_chaild]
                    list_attribute_name = self.dict_atribut_level_3[
                        self.dict_class_content[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]]]
                    file.write("<h3 align=left style=margin-left:70 id="+element_chaild+">" + element_index + " " + element[1] + "</h3>")
                    for tmp_attr in element[3]:
                        if tmp_attr["f4"] != "r" and tmp_attr["f2"] is not None and tmp_attr["f2"] != "":
                            for attr in list_attribute_name:
                                if tmp_attr["f1"] in attr:
                                    if attr[4] == "s" :
                                        file.write("<p align=left style=margin-left:90><b>" + attr[1] + "</b> : " + tmp_attr["f2"] + "</p>")
                                    elif tmp_attr["f1"]!=self.id_attr_index:
                                        file.write("<h4 align=left style=margin-left:90>" + attr[1] + "</h4>")
                                        file.write("<div align=left style=margin-left:90>" + tmp_attr["f2"] + "</div>")
                        if element_chaild != "":
                           self.recursive_for_report_html(dict_level3, v, level + 2, file)
        return (element_chaild)  # вернули id элемента дочернего
    def init_tree(self):

        #к этому моменту необходимо иметь словари, заполненные данными для построения дерева методов.

        self.ui_Simma.treeWidget_SIMMA.clear()
        tw_main = self.ui_Simma.treeWidget_SIMMA
        tw_main.setHeaderLabels([".."])
        tw = QtWidgets.QTreeWidgetItem(tw_main, [self.ui_Simma.comboBox_report.currentText()])
        tw.setText(0, self.ui_Simma.comboBox_report.currentText())
        for tmp in self.dict_list_level1_2[self.ui_Simma.comboBox_report.currentText()]:
            if isinstance(tmp,str):
               sg = QtWidgets.QTreeWidgetItem(tw, [self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp][1]])
               sg.setText(0, self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp][1])
               i=1
               child=self.recursive(self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()],
                                     self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp],i,sg)

            else:
               tmp_item=list(tmp.keys())[0]
               sg = QtWidgets.QTreeWidgetItem(tw, [self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item][1]])
               sg.setText(0, self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item][1])
               child_list=tmp[tmp_item]
               for tmp_child in child_list:
                   ss = QtWidgets.QTreeWidgetItem(sg)
                   ss.setText(0, self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_child][1])
                   sg.addChild(ss)
                   i = 3
                   child = self.recursive(self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()],
                                          self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_child],i,ss)
            self.ui_Simma.treeWidget_SIMMA.setCurrentItem(tw)
            for k,v in self.dict_element_level_1.items():
                if v[1]== self.ui_Simma.comboBox_report.currentText():
                   for tmp_element in self.dict_element_level_1[k][3]:
                       if tmp_element["f1"] ==self.id_attr_text_list_1:
                          text_1=tmp_element["f2"]
            self.ui_Simma.textEdit_item.setText(text_1)
    def contextMenuEvent(self, point):
        menu = QtWidgets.QMenu()
        add_element = menu.addAction("Создать документ по разделу")
        add_element.triggered.connect(self.slice_doc_html)
        menu.exec(self.ui_Simma.treeWidget_SIMMA.mapToGlobal(point))
    def all_doc_docx(self):
        # печатаем весь документ. начинаем с вывода данных содержащихся в классе Отчет
        doc = None
        doc = docx.Document()
        doc.sections.different_first_page_header_footer = True
        for k,v in self.dict_element_level_1.items():
            if v[1]== self.ui_Simma.comboBox_report.currentText():
                for tmp_element in self.dict_element_level_1[k][3]:
                    if tmp_element["f1"]== self.id_attr_colontitul_list_1_up:
                       header_1=tmp_element["f2"]
                    elif tmp_element["f1"]==self.id_attr_colontitul_list_1_down:
                       footer_1 = tmp_element["f2"]
                    elif tmp_element["f1"] == self.id_attr_colontitul_list_all_up:
                       header_all = tmp_element["f2"]
                    elif tmp_element["f1"] ==self.id_attr_colontitul_list_all_down:
                       footer_all = tmp_element["f2"]
                    elif tmp_element["f1"] ==self.id_attr_text_list_1:
                       text_1=tmp_element["f2"]
                header1 = doc.sections[0].first_page_header.paragraphs[0]  # добавляем верхний колонтитул первой страницы
                footer1 = doc.sections[0].first_page_footer.paragraphs[0]  # доступ к нижнему колонтитулу первой страницы
                # заполняем колонититулы первой страницы
                soup = BeautifulSoup(header_1, 'html.parser')
                header1.add_run(soup.get_text())
                soup = BeautifulSoup(footer_1, 'html.parser')
                footer1.add_run(soup.get_text())
                soup = BeautifulSoup(text_1, 'html.parser')
                # выводим текст первой страницы
                doc.add_paragraph(soup.get_text())
                # Добавим разрыв страницы
                doc.add_page_break()
                # создаем колонтитулы для остальных листов
                header = doc.sections[0].header.paragraphs[0]
                footer = doc.sections[0].footer.paragraphs[0]
                soup = BeautifulSoup(header_all, 'html.parser')
                header.add_run(soup.get_text())
                soup = BeautifulSoup(footer_all, 'html.parser')
                footer.add_run(soup.get_text())
        # приступаем к созданию огравления
        style = doc.styles['Normal']
        # название шрифта
        style.font.name = 'Calibri'
        # размер шрифта
        p = doc.add_paragraph()
        run=p.add_run("Оглавление")
        run.font.bold = True
        run.font.size = Pt(22)
        #формирование тела документа аналогично дереву.
        for tmp in self.dict_list_level1_2[self.ui_Simma.comboBox_report.currentText()]:
            if isinstance(tmp, str):
                element_print=self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp]
                p1 = doc.add_paragraph()
                run=p1.add_run(element_print[1])
                run.font.bold = True
                run.font.size = Pt(18)
                p2 = doc.add_paragraph()
                for tmp_print in element_print[3]:
                    if tmp_print["f1"]==self.id_attr_level2_annotation:
                        soup = BeautifulSoup(tmp_print["f2"], 'html.parser')
                        run=p2.add_run(soup.get_text())
                        run.font.bold = False
                        run.font.size = Pt(12)
                if element_print[2] is not None and element_print[2]!="":
                    p3 = doc.add_paragraph()
                    run = p3.add_run("(")
                    soup = BeautifulSoup(element_print[2])
                    run = p3.add_run(soup.get_text())
                    run = p3.add_run(")")
                    run.font.italic=True
                i = 1
                self.recursive_for_report(self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()],
                                       self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp], i, doc)
            else:
                tmp_item = list(tmp.keys())[0]
                element_print = self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item]
                p1 = doc.add_paragraph()
                run = p1.add_run(element_print[1])
                run.font.bold = True
                run.font.size = Pt(18)
                p2 = doc.add_paragraph()
                for tmp_print in element_print[3]:
                    if tmp_print["f1"] == self.id_attr_level2_annotation:
                        soup = BeautifulSoup(tmp_print["f2"], 'html.parser')
                        run = p2.add_run(soup.get_text())
                        run.font.bold = False
                        run.font.size = Pt(12)
                if element_print[2] is not None and element_print[2] != "":
                    p3 = doc.add_paragraph()
                    run = p3.add_run("(")
                    soup = BeautifulSoup(element_print[2])
                    run = p3.add_run(soup.get_text())
                    run = p3.add_run(")")
                    run.font.italic = True
                child_list = tmp[tmp_item]
                for tmp_child in child_list:
                    element_print = self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_child]
                    p1 = doc.add_paragraph()
                    run = p1.add_run(element_print[1])
                    run.font.bold = True
                    run.font.size = Pt(16)
                    p2 = doc.add_paragraph()
                    for tmp_print in element_print[3]:
                        if tmp_print["f1"] == self.id_attr_level2_annotation:
                            soup = BeautifulSoup(tmp_print["f2"], 'html.parser')
                            run = p2.add_run(soup.get_text())
                            run.font.bold = False
                            run.font.size = Pt(12)
                    if element_print[2] is not None and element_print[2] != "":
                        p3 = doc.add_paragraph()
                        run = p3.add_run("(")
                        soup = BeautifulSoup(element_print[2])
                        run = p3.add_run(soup.get_text())
                        run = p3.add_run(")")
                        run.font.italic = True
                    i = 3
                    self.recursive_for_report(
                        self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()],
                        self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_child], i, doc)
        doc.save(self.ui_Simma.comboBox_report.currentText()+'.docx')
    def all_doc_html(self):
        # создаем весь документ. начинаем с вывода данных содержащихся в классе Отчет
        self.ui_Simma.progressBar_report.maximum = 100
        valBar=0
        self.ui_Simma.progressBar_report.setValue(valBar)
        filename=self.ui_Simma.comboBox_report.currentText()
        filename = re.sub(r"[^A-Za-zа-яА-Я]+", '_', filename)
        file = open(filename + '.html', 'w', encoding='utf-16')
        file.write("<html>")
        for k,v in self.dict_element_level_1.items():
            if v[1]== self.ui_Simma.comboBox_report.currentText():
                for tmp_element in self.dict_element_level_1[k][3]:
                    if tmp_element["f1"]== self.id_attr_colontitul_list_1_up:
                       header_1=tmp_element["f2"]
                    elif tmp_element["f1"]==self.id_attr_colontitul_list_1_down:
                       footer_1 = tmp_element["f2"]
                    elif tmp_element["f1"] == self.id_attr_colontitul_list_all_up:
                       header_all = tmp_element["f2"]

                    elif tmp_element["f1"] ==self.id_attr_text_list_1:
                       text_1=tmp_element["f2"]
                # заполняем колонититулы первой страницы
                file.write("<div align=left> " + header_1 + "</div>")
                file.write("<br>")
                # выводим текст первой страницы
                file.write("<div align=center> "+text_1+"</div>")
                file.write("<br>")
                file.write("<br>")
                file.write("<br>")
                file.write("<br>")
                file.write("<br>")
                file.write("<div align=center> " + footer_1 + "</div>")
                file.write("<br>")
                file.write("<br>")
                footer_all="Создано из СиММА. Дата : "+str(date.today())
                file.write("<div align=left> " + footer_all + "</div>")
                user="Пользователь : " + self.login
                file.write("<div align=left> " + user + "</div>")
                # Добавим разрыв страницы
                file.write("<hr>")
                # приступаем к созданию огравления
                file.write("<h1 align=center> Оглавление </h1>")
                list_for_section=self.dict_list_level1_2[self.ui_Simma.comboBox_report.currentText()]
                for tmp in list_for_section:
                    if isinstance(tmp, str):
                       tmp_text=self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp] + " " + self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp][1]
                       file.write("<h2 align=LEFT><a href=#"+self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp][0]+">" + tmp_text + "</a></h2>")
                    if isinstance(tmp, dict):
                       tmp_item = list(tmp.keys())[0]
                       tmp_text = self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_item] + " " + \
                                  self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item][1]
                       file.write("<h2 align=LEFT><a href=#"+self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item][0]+">"+ tmp_text + "</a></h2>")
                       list_values=list(tmp.values())
                       for tmp_in_dict in list_values[0]:
                           tmp_text = self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_in_dict] + " " + \
                                      self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_in_dict][1]
                           file.write("<h2 align=LEFT><a href=#"+self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_in_dict][0]+">" + tmp_text + "</a></h2>")
                file.write("<br>")
                file.write("<br>")
                file.write("<br>")
                # Закончили оглавление - Добавим разрыв страницы
                file.write("<hr>")
        #формирование тела документа аналогично дереву.
        for tmp in self.dict_list_level1_2[self.ui_Simma.comboBox_report.currentText()]:
            if isinstance(tmp, str):
                tmp_text = self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp] + " " + \
                           self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp][1]
                file.write("<h2 align=LEFT id="+ tmp +">" + tmp_text + "</h2>")
                element_print=self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp]
                for tmp_print in element_print[3]:
                    if tmp_print["f1"]==self.id_attr_level2_preambula:
                       file.write("<h4 align=left>"+tmp_print["f2"]+" </h4>")
                i = 1
                dict_level3=self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()]
                dict_for_sort={}
                for k,v in dict_level3.items():
                    if v is not None and v!="":
                        dict_for_sort[k]=v
                marklist = sorted(dict_for_sort.items(), key=lambda x: x[1])
                self.recursive_for_report_html(dict(marklist),self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp], i, file)
            else:
                tmp_item = list(tmp.keys())[0]
                element_print = self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item]
                tmp_text = self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_item] + " " + \
                           self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_item][1]
                file.write("<h2 align=LEFT id="+ tmp_item +">" + tmp_text + "</h2>")
                for tmp_print in element_print[3]:
                    if tmp_print["f1"] == self.id_attr_level2_preambula:
                        file.write("<h4 align=left>"+tmp_print["f2"]+" </h4>")

                child_list = tmp[tmp_item]
                for tmp_child in child_list:
                    element_print = self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_child]
                    tmp_text = self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_child] + " " + \
                               self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_child][1]
                    file.write("<h2 align=LEFT  style=margin-left:50 id="+self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_child][0]+">" + tmp_text + "</h2>")
                    for tmp_print in element_print[3]:
                        if tmp_print["f1"] == self.id_attr_level2_preambula:
                            file.write("<h4 align=left>"+tmp_print["f2"]+" </h4>")
                    if self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()][tmp_child][1] == "Перечень методов":
                       self.all_method(file)
                    i = 3
                    dict_level3 = self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()]
                    dict_for_sort = {}
                    for k, v in dict_level3.items():
                        if v is not None and v != "":
                            dict_for_sort[k] = v
                    marklist = sorted(dict_for_sort.items(), key=lambda x: x[1])
                    self.recursive_for_report_html(dict(marklist),self.dict_element_index[self.ui_Simma.comboBox_report.currentText()][tmp_child], i, file)
        file.write("<table border=1>")
        file.write("<colgroup><col  span = \"2\"  style = \"background:Khaki\"><col span = \"2\" style = \"background-color:LightCyan\"></colgroup>")
        file.write("<tr><th width=\"50\"> № п/п        </th><th width=\"450\">Наименование контента              </th>><th width=\"150\">Дата изменения                </th><th>Автор изменения</th></tr>")
        i=0
        valBar=10
        self.ui_Simma.progressBar_report.setValue(valBar)
        j=90/len(marklist)
        for element in marklist:
            i+=1
            tmp_element=All_Class.Simma_element
            tmp_element.e_id=element[0]
            about_element=tmp_element.gea(tmp_element,self.url,self.mid,self.sid)
            for tmp_attr in about_element:
                if tmp_attr[0]=="<Дата изменения>":
                   date_correct=tmp_attr[1]
                elif tmp_attr[0]=="Название":
                   name_attr=tmp_attr[1]
                elif tmp_attr[0] == "<Внёс изменения>":
                    name_autor= tmp_attr[1]
            about_history=tmp_element
            file.write("<tr><td>"+str(i)+"</td><td>"+name_attr+"</td><td>"+str(date_correct)+"</td><td>"+name_autor+"</td></tr>")
            valBar+=round(j)
            self.ui_Simma.progressBar_report.setValue(valBar)
        file.write("</table>")
        file.write("</html>")
        file.close()
        valBar = 100
        self.ui_Simma.progressBar_report.setValue(valBar)
        self.show_Ok_messagebox(1, ".html")
    def all_method(self,file):
        dict_level3 = self.dict_element_index3[self.ui_Simma.comboBox_report.currentText()]
        dict_for_sort = {}
        for k, v in dict_level3.items():
            if v is not None and v != "":
                dict_for_sort[k] = v
        marklist = sorted(dict_for_sort.items(), key=lambda x: x[1])
        i=0
        file.write("<table border=1>")
        file.write("<colgroup><col  span = \"2\" style = \"background:Khaki\"></colgroup>")
        file.write("<tr><th width=\"50\"> № п/п        </th><th width=\"650\">Наименование метода              </th>></tr>")
        for tmp_method in marklist:
            i+=1
            name_method=self.dict_element_level_3[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]][tmp_method[0]][1]
            descr_method=self.dict_element_level_3[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]][tmp_method[0]][2]
            file.write("<tr><td>"+str(i)+"</td><td><a href=#"+tmp_method[0]+">"+name_method+"</a></td></tr>")
        file.write("</table>")
    def slice_doc_html(self):
        valBar = 0
        self.ui_Simma.progressBar_report.setValue(valBar)
        filename=self.ui_Simma.comboBox_report.currentText() + '_' + self.ui_Simma.treeWidget_SIMMA.currentItem().text(0)
        filename=re.sub(r"[^A-Za-zа-яА-Я]+",'_', filename)
        file = open(filename+ '.html', 'w', encoding='utf-16')
        file.write("<html>")
        valBar = 5
        self.ui_Simma.progressBar_report.setValue(valBar)
        name_element = self.ui_Simma.treeWidget_SIMMA.currentItem().text(0)
        list_for_label = ""
        list_attribute_name = self.dict_atribut_level_3[
            self.dict_class_content[self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]]]
        for k, v in self.dict_element_level_3[
            self.dict_item_content[self.ui_Simma.comboBox_report.currentText()]].items():
            if v[1] == name_element:
                for tmp_attr in v[3]:
                    if tmp_attr["f4"] != "r" and tmp_attr["f2"] is not None and tmp_attr["f2"] != "":
                        for attr in list_attribute_name:
                            if tmp_attr["f1"] in attr:
                               if tmp_attr["f1"] != self.id_attr_index:
                                    list_for_label = list_for_label + (
                                                "<p align=left><b>" + attr[1] + "</b> : " + tmp_attr["f2"] + "</p>")
       #                         elif tmp_attr["f1"] != self.id_attr_index:
       #                             list_for_label = list_for_label + ("<h4 align=left>" + attr[1] + "</h4>")
       #                             list_for_label = list_for_label + ("<div align=left>" + tmp_attr["f2"] + "</div>")
        for key, val in self.dict_class.items():
            if val == "Раздел отчета":
                list_attribute_name = self.dict_atribut_level_2[key]
        valBar = 15
        self.ui_Simma.progressBar_report.setValue(valBar)
        for k, v in self.dict_element_level_2[self.ui_Simma.comboBox_report.currentText()].items():
            if v[1] == name_element:
                for tmp_attr in v[3]:
                    if tmp_attr["f1"] == self.id_attr_level2_preambula:
                        if tmp_attr["f2"] is not None and tmp_attr["f2"] != "":
                            list_for_label = list_for_label + ("<h4 align=left> Преамбула раздела </h4>")
                            list_for_label = list_for_label + ("<div align=left>" + tmp_attr["f2"] + "</div>")
        list_for_label = "<H2>" + name_element + "</H2>" + list_for_label
        list_for_label=list_for_label+"</html>"
        file.write(list_for_label)
        file.close()
        valBar = 100
        self.ui_Simma.progressBar_report.setValue(valBar)
        self.show_Ok_messagebox(3, ".html")
    def init_graph(self):
        if len(self.sid) == 32:
            self.ui_Simma.listWidget_report.clear()
            self.ui_Simma.listWidget_attr_report.clear()
            sc = All_Class.Simma_class
            sc.o_id = list(self.dict_class.keys())[0]
            about_class = sc.gtvt(sc, self.url, self.mid, self.sid)
            for tmp in about_class:
                self.ui_Simma.listWidget_report.addItem(tmp[1])
            attr_diagr=["Автор контента","Автор изменений контента","Статусы контента"]
            for tmp in  attr_diagr:
                self.ui_Simma.listWidget_attr_report.addItem(tmp)
            self.ui_Simma.listWidget_report.setCurrentRow(0)
            self.ui_Simma.listWidget_attr_report.setCurrentRow(0)
    def def_diagramm(self):
        match self.ui_Simma.listWidget_attr_report.currentItem().text():
            case "Автор контента":
                 self.view_round_diagramm_standart(self.ui_Simma.listWidget_report.currentItem().text(),self.ui_Simma.listWidget_attr_report.currentItem().text())
            case "Автор изменений контента":
                 self.view_round_diagramm_standart(self.ui_Simma.listWidget_report.currentItem().text(),self.ui_Simma.listWidget_attr_report.currentItem().text())
            case "Статусы контента":
                 self.view_bar_diagramm_one_attr(self.ui_Simma.listWidget_report.currentItem().text(), self.id_attr_status)
    def view_round_diagramm_standart(self, report, type_report):
        klass=self.dict_class_content[self.dict_item_content[report]]
        sc = All_Class.Simma_class
        sc.o_id = klass
        about_class = sc.gca(sc, self.url, self.mid, self.sid)
        list_for_sevt = []
        for tmp in about_class:
            list_for_sevt.append(tmp[0])
        list_for_sevt.append("descr")
        list_for_sevt.append("syscr")
        sc.sevt(sc, self.url, self.mid, self.sid, list_for_sevt, [])
        about_class = sc.gtvt(sc, self.url, self.mid, self.sid)
        sc.sevt(sc, self.url, self.mid, self.sid, ['name'], [])
        tmp_series = []
        match type_report:
            case "Автор контента":
                 attr_for_select=0
            case "Автор изменений контента":
                 attr_for_select = 2
        for tmp_element in about_class:
            for tmp_attr in tmp_element[3]:
                if tmp_attr["f1"] =="syscr":
                   if tmp_attr["f3"][attr_for_select] in tmp_series:
                      index=tmp_series.index(tmp_attr["f3"][attr_for_select])
                      tmp_series[index+1]=tmp_series[index+1]+1
                   else:
                      tmp_series.append(tmp_attr["f3"][attr_for_select])
                      tmp_series.append(1)
        series=QPieSeries()
        for i in range(0,len(tmp_series)-1,2):
            series.append(tmp_series[i],tmp_series[i+1])
        series.setLabelsVisible(True)
#        series.setLabelsPosition(QtChart.QPieSlice.LabelInsideHorizontal)
        for slice in series.slices():
            slice.setLabel("{:.2f}%".format(100 * slice.percentage()))
        chart = QChart()
        chart.addSeries(series)
        chart.createDefaultAxes()
        chart.setTheme(QChart.ChartTheme.ChartThemeBlueCerulean)
        chart.setAnimationOptions(QChart.AnimationOption.AllAnimations)
        chart.setTitle(self.ui_Simma.listWidget_attr_report.currentItem().text()+" ( " +self.ui_Simma.listWidget_report.currentItem().text()+" )")
        chart.legend().setVisible(True)
        chart.legend().setAlignment(Qt.AlignBottom)
        for i in range(round(len(tmp_series)/2)):
            if i==0:
               chart.legend().markers(series)[i].setLabel(tmp_series[i])
            else:
               chart.legend().markers(series)[i].setLabel(tmp_series[i*2])
        my_slice=series.slices()[0]
        my_slice.setExploded(True)
        my_slice.setLabelVisible(True)
        my_slice.setBrush(Qt.GlobalColor.red)
        chart.legend().hide()
        chart.legend().setVisible(True)
        chart.legend().setAlignment(Qt.AlignmentFlag.AlignBottom)
        self.ui_Simma.graphicsView_diagram.setChart(chart)
    def view_bar_diagramm_one_attr(self,report,id_attr):
        klass = self.dict_class_content[self.dict_item_content[report]]
        sc = All_Class.Simma_class
        sc.o_id = klass
        about_class = sc.gca(sc, self.url, self.mid, self.sid)
        list_for_sevt = []
        for tmp in about_class:
            list_for_sevt.append(tmp[0])
        sc.sevt(sc, self.url, self.mid, self.sid, list_for_sevt, [])
        about_class = sc.gtvt(sc, self.url, self.mid, self.sid)
        sc.sevt(sc, self.url, self.mid, self.sid, ['name'], [])
        tmp_series = []
        for tmp_element in about_class:
            for tmp_attr in tmp_element[3]:
                if tmp_attr["f1"] ==id_attr:
                     if tmp_attr["f2"] in tmp_series:
                         index=tmp_series.index(tmp_attr["f2"])
                         tmp_series[index+1]=tmp_series[index+1]+1
                     else:
                         tmp_series.append(tmp_attr["f2"])
                         tmp_series.append(1)
        set0 = QBarSet("Статусы элементов контента")
        set0.clicked().self.parametr=2
        set0.clicked.connect(self.chaild_window)

        for i in range(1,len(tmp_series),2):
            set0.append(tmp_series[i])
        series = QStackedBarSeries()
        series.append(set0)
        chart = QChart()
        chart.addSeries(series)
        chart.setTitle("Количество элементов (шт.) в соответсвующе статусе")
        chart.setAnimationOptions(QChart.AnimationOption.SeriesAnimations)
        categories = []
        for i in range(0,len(tmp_series) - 1,2):
            categories.append(tmp_series[i])
        axisX = QBarCategoryAxis()
        axisX.append(categories)
        chart.addAxis(axisX, Qt.AlignmentFlag.AlignBottom)
        series.attachAxis(axisX)
        axisY = QValueAxis()
        chart.addAxis(axisY, Qt.AlignmentFlag.AlignLeft)
        series.attachAxis(axisY)
        chart.legend().setVisible(True)
        chart.legend().setAlignment(Qt.AlignmentFlag.AlignBottom)
        self.ui_Simma.graphicsView_diagram.setChart(chart)

    def chaild_window(self):
        print(self.parametr)

if __name__ == '__main__':
    app =QApplication(sys.argv)
    mainUi = ui_Simma()
    app.exec()

